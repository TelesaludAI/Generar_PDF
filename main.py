from fastapi import FastAPI, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import os
import aiofiles

# === CONFIGURACIÓN ===
SERVICE_ACCOUNT_FILE = "credentials.json"  # archivo JSON de la cuenta de servicio
FOLDER_ID = "1E4brfQPnwurTivj8mjpdAtc_ByprLbNL"    # ID de la carpeta de destino en Drive

SCOPES = ["https://www.googleapis.com/auth/drive.file"]

app = FastAPI(title="Generador de documentos en Drive")

# === MODELO DE ENTRADA ===
class DocumentoData(BaseModel):
    titulo: str
    contenido: str
    autor: str | None = None

# === FUNCIÓN AUXILIAR ===
def upload_to_drive(filepath: str, filename: str):
    """Sube un archivo a Google Drive usando una cuenta de servicio"""
    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES
    )
    service = build("drive", "v3", credentials=creds)

    file_metadata = {"name": filename, "parents": [FOLDER_ID]}
    media = MediaFileUpload(filepath, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    file = service.files().create(
        body=file_metadata,
        media_body=media,
        fields="id, webViewLink"
    ).execute()

    return file.get("id"), file.get("webViewLink")

# === ENDPOINT PRINCIPAL ===
@app.post("/")
async def crear_documento(data: DocumentoData):
    try:
        # Nombre temporal del archivo
        filename = f"{data.titulo.replace(' ', '_')}.docx"
        
        # Crear el documento DOCX
        doc = Document()
        doc.add_heading(data.titulo, level=1)
        doc.add_paragraph(data.contenido)
        if data.autor:
            doc.add_paragraph(f"Autor: {data.autor}")
        doc.save(filename)

        # Subir a Google Drive
        file_id, drive_link = upload_to_drive(filename, filename)

        # Eliminar archivo local
        os.remove(filename)

        return JSONResponse({
            "message": "✅ Documento creado y subido a Google Drive",
            "drive_file_id": file_id,
            "drive_link": drive_link
        })

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al crear o subir el documento: {str(e)}")

